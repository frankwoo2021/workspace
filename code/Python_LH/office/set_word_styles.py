#########################################################################
# 作者:李宁（蒙娜丽宁），UnityMarvel创始人
#
# 微信公众号：极客起源
#
# B站：https://space.bilibili.com/477001733
#
# Copyright © 2022 Lining. All rights reserved.
#
# 版本: 2.0
#########################################################################

import docx

doc = docx.Document('demo.docx')
doc.paragraphs[0].style = 'Normal'

doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save('new.docx')
