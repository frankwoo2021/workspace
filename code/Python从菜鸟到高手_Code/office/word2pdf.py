#########################################################################
# 作者:李宁（蒙娜丽宁），UnityMarvel创始人
#
# 微信公众号：极客起源
#
# B站：https://space.bilibili.com/477001733
#
# Copyright © 2022 Lining. All rights reserved.
#
# 版本: 2.0
#########################################################################


import win32com.client
import docx
wordFilename = 'word_document.docx'
pdfFilename = 'pdf_filename.pdf'

doc = docx.Document()
doc.add_picture('book.png', width=docx.shared.Inches(1.6),
                            height=docx.shared.Cm(4))
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)

p = doc.add_paragraph('Word2Pdf')
p.runs[0].underline = True
p.runs[0].italic = True
doc.save(wordFilename)
import os



wordPath = os.path.abspath(os.path.dirname(wordFilename)) + os.sep + wordFilename
pdfPath =  os.path.abspath(os.path.dirname(pdfFilename)) + os.sep + pdfFilename


wdFormatPDF = 17
wordObj = win32com.client.Dispatch('Word.Application')
docObj = wordObj.Documents.Open(wordPath)
docObj.SaveAs(pdfPath, FileFormat=wdFormatPDF)
docObj.Close()
wordObj.Quit()